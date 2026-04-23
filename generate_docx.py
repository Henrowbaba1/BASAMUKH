from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_profile():
    doc = Document()

    # --- COVER PAGE ---
    # Centered Header
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27) # A4

    # Add Cover Image (if exists)
    if os.path.exists('cover.png'):
        doc.add_picture('cover.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- MISSION & VISION ---
    doc.add_heading('Basamukh Seed Integrated Services', 0)
    
    p = doc.add_paragraph('Cultivating the Future of African Agriculture')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(27, 94, 32)
    run.italic = True

    doc.add_heading('Company Overview', level=1)
    doc.add_paragraph(
        'Basamukh Seed Integrated Services is a premier agricultural enterprise based in Samaru Zaria, Kaduna. '
        'We specialize in the production, processing, and distribution of high-quality seeds that drive agricultural productivity across Africa.'
    )

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Our Vision'
    hdr_cells[1].text = 'Our Mission'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'To become a leading seed production and agricultural solutions company in Nigeria and across Africa, contributing to food security and sustainable farming systems.'
    row_cells[1].text = 'To provide farmers with high-quality, affordable, and improved seeds that enhance productivity, profitability, and sustainability in agriculture.'

    doc.add_heading('Core Values', level=2)
    values = ['Quality Assurance', 'Integrity & Transparency', 'Innovation in Agriculture', 'Customer Satisfaction', 'Sustainability']
    for val in values:
        doc.add_paragraph(val, style='List Bullet')

    doc.add_page_break()

    # --- THE SEED JOURNEY (INFOGRAPHIC) ---
    doc.add_heading('The Seed Quality Journey', level=1)
    doc.add_paragraph(
        'Our process ensures that every seed delivered to a farmer is of the highest genetic and physical purity. '
        'From rigorous R&D to final distribution, we maintain global standards of excellence.'
    )
    
    if os.path.exists('infographic.png'):
        doc.add_picture('infographic.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Our Seed Portfolio', level=1)
    seeds = {
        'Maize': 'High-yield hybrids and open-pollinated varieties developed for drought tolerance.',
        'Sorghum': 'Climate-smart varieties with improved nutritional profiles.',
        'Cowpea': 'Quick-maturing, insect-resistant varieties.',
        'Groundnut': 'High-oil content seeds with excellent disease resistance.',
        'Sesame': 'Export-quality seeds for international markets.',
        'Soybean': 'Protein-rich seeds with superior germination rates.'
    }
    
    for seed, desc in seeds.items():
        p = doc.add_paragraph()
        run = p.add_run(f'{seed}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # --- CONTACT ---
    doc.add_heading('Contact & Reach', level=1)
    doc.add_paragraph('Basamukh Seed Integrated Services is headquartered in the agricultural heart of Northern Nigeria.')
    
    contact_table = doc.add_table(rows=3, cols=2)
    contact_table.cell(0, 0).text = '📍 Address'
    contact_table.cell(0, 1).text = 'No. 10 Tsamiya Road, Hayin Dogo Samaru - Zaria, Kaduna State, Nigeria'
    
    contact_table.cell(1, 0).text = '📞 Phone'
    contact_table.cell(1, 1).text = '07019800030 | 08065307616'
    
    contact_table.cell(2, 0).text = '✉️ Email'
    contact_table.cell(2, 1).text = 'basamukhseed@gmail.com'

    # Save the document
    doc.save('Basamukh_Seed_Integrated_Services_Profile.docx')
    print("Document saved successfully as Basamukh_Seed_Integrated_Services_Profile.docx")

if __name__ == "__main__":
    create_profile()
